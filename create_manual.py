from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Tiêu đề chính
title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG HỆ THỐNG', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('QUẢN LÝ KHOA HỌC CÔNG NGHỆ VÀ CHUYỂN ĐỔI SỐ - EVNGENCO2').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# Phần 1: Giới thiệu chung
h1 = doc.add_heading('1. Giới thiệu chung', level=1)
doc.add_paragraph('Hệ thống Quản lý Khoa học Công nghệ và Chuyển đổi số (EVNGENCO2) là phần mềm giúp các cán bộ nhân viên đăng ký sáng kiến, đề tài khoa học công nghệ, đồng thời hỗ trợ Hội đồng xét duyệt thẩm định, chấm điểm và ra quyết định một cách minh bạch, tuân thủ theo các quy định hiện hành (Quyết định 722/QĐ-EVN, Nghị quyết 388/NQ-HĐQT).')

# Phần 2: Đăng nhập hệ thống
doc.add_heading('2. Đăng nhập hệ thống', level=1)
doc.add_paragraph('Truy cập vào đường dẫn của hệ thống (ví dụ: https://genco2-xyz.onrender.com).')
p2 = doc.add_paragraph()
p2.add_run('Các cấp tài khoản:').bold = True
doc.add_paragraph('Tài khoản Admin: Dành cho Quản trị viên quản lý danh sách tài khoản, cài đặt hệ thống.', style='List Bullet')
doc.add_paragraph('Tài khoản Đăng ký (Nhân viên): Dành cho tác giả nộp sáng kiến.', style='List Bullet')
doc.add_paragraph('Tài khoản Hội đồng (Ban Lãnh đạo): Dành cho thành viên Hội đồng thẩm định và Tổng Giám đốc.', style='List Bullet')

# Phần 3: Chức năng dành cho Người đăng ký sáng kiến
doc.add_heading('3. Chức năng dành cho Người đăng ký sáng kiến', level=1)
doc.add_paragraph('Sau khi đăng nhập bằng tài khoản Nhân viên, người dùng sẽ được chuyển đến trang Biểu mẫu Đăng ký sáng kiến.')
doc.add_paragraph('Bước 1: Điền thông tin chung (Tên sáng kiến, lĩnh vực, tác giả chủ trì, thành viên phối hợp).', style='List Number')
doc.add_paragraph('Bước 2: Điền nội dung cốt lõi, nêu rõ tính mới và thực trạng (Đối chiếu với các quy định cũ như Quyết định 93/QĐ-EVNGENCO2).', style='List Number')
doc.add_paragraph('Bước 3: Tải lên các tệp tin minh chứng (hỗ trợ PDF, DOCX, XLSX, CAD) với dung lượng tối đa 50MB.', style='List Number')
doc.add_paragraph('Bước 4: Tích chọn Cam kết bản quyền và nhấn "Gửi hồ sơ".', style='List Number')

# Phần 4: Chức năng dành cho Hội đồng Xét duyệt
doc.add_heading('4. Chức năng dành cho Hội đồng Xét duyệt (Ban Lãnh đạo)', level=1)
doc.add_paragraph('Đăng nhập bằng tài khoản Hội đồng để vào Bảng điều khiển (Dashboard).')
p4 = doc.add_paragraph()
p4.add_run('Các tính năng chính:').bold = True
doc.add_paragraph('Bảng Thống kê: Xem tổng quan số lượng hồ sơ (Tổng số, Chờ duyệt, Đang thẩm định, Đã duyệt).', style='List Bullet')
doc.add_paragraph('Thẩm định hồ sơ: Nhấp vào nút "Thẩm định" tại hồ sơ cần xét duyệt. Hệ thống sẽ mở bảng chấm điểm chi tiết (Tính mới, Khả thi, Hiệu quả).', style='List Bullet')
doc.add_paragraph('Ý kiến Lãnh đạo: Ghi nhận ý kiến chỉ đạo trực tiếp của Tổng Giám đốc dựa trên đề xuất của Hội đồng chuyên môn.', style='List Bullet')
doc.add_paragraph('Quyết định & Kết xuất: Sử dụng tính năng "Ký số Phê duyệt (TGĐ)" hoặc "Tạo dự thảo Quyết định" để tự động sinh văn bản có tiêu ngữ của EVNGENCO2.', style='List Bullet')

# Phần 5: Tính năng bổ trợ
doc.add_heading('5. Tính năng bổ trợ', level=1)
doc.add_paragraph('Tra cứu văn bản pháp lý: Giao diện Hội đồng cung cấp mục Tra cứu nhanh các văn bản liên quan (QĐ 722/QĐ-EVN, NQ 388/NQ-HĐQT) để đối chiếu nhanh chóng.', style='List Bullet')
doc.add_paragraph('Kết xuất báo cáo: Hỗ trợ xuất Biên bản họp Hội đồng và Báo cáo thống kê ra file Word.', style='List Bullet')

# Lưu file
file_name = 'HuongDanSuDung_EVNGENCO2.docx'
doc.save(file_name)
print(f'Saved {file_name}')
